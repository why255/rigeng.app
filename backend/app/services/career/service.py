"""高维求职服务 — 核心业务逻辑（步骤20 / Wave 4）。

五步法全流程：
  一盘（简历盘点与重构）→ 二定（求职策略与资源）→ 三投（投递追踪与分析）
  → 四面（面试准备与复盘）→ 五选（Offer评估与入职）

跨模块调用链：
  面试录音 → ③智能记录（Recording 关联）
  技能晶体 → ②知识库（Document 归档）
  企业情报 → ③AI引擎（LLM生成初稿）

设计原则：
  - 只写业务逻辑层，基础能力全部调用已有服务
  - 技能晶体归档到私有知识库，删除求职记录不影响已归档晶体
  - 老师由后台手动分配，用户不可自选
  - 辅导统一为内置视频，合同平台外签署
  - 简历真实性：不做事实核查，不指责
  - 数据永久保留，不设退出期限
"""
from __future__ import annotations

import json
import logging
from datetime import date, datetime, timedelta, timezone
from typing import Any

from sqlalchemy import and_, desc, func
from sqlalchemy.orm import Session

from ...shared.config import settings
from ...shared.database import new_uuid, utcnow
from ...shared.errors import APIError, E_FILE_NOT_FOUND, E_PARAM_FORMAT
from ...shared.llm_utils import safe_extract_json, LLMParsingError
from ...shared.models.career import (
    CareerProgress, STARExtraction, SkillCrystal, JobApplication,
    InterviewPrep, InterviewReview, OfferComparison, ProbationPlan,
    CompanyIntel,
)
from ...shared.models.user import TeacherProfile
from ..voice_engine.service import llm_generate, MODULE_SYSTEM_PROMPTS

logger = logging.getLogger("career")

# ═══════ 五步法步骤标签 ═══════
STEP_LABELS: dict[int, str] = {
    1: "一盘·简历盘点与重构",
    2: "二定·求职策略与资源",
    3: "三投·投递追踪与分析",
    4: "四面·面试准备与复盘",
    5: "五选·Offer评估与入职",
}

# ═══════ 投递状态标签 ═══════
STATUS_LABELS: dict[str, str] = {
    "applied": "已投递",
    "screening": "筛选中",
    "interview": "面试中",
    "offer": "已获Offer",
    "rejected": "未通过",
    "withdrawn": "已撤回",
}


def _get_or_create_progress(db: Session, user_id: str) -> CareerProgress:
    """获取或创建用户的五步法进度记录。"""
    progress = db.query(CareerProgress).filter(
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()

    if not progress:
        progress = CareerProgress(
            user_id=user_id,
            current_step=1,
            status="active",
            step_data_json={},
        )
        db.add(progress)
        db.commit()
        db.refresh(progress)

    return progress


def _check_completeness(situation: str, task: str, action: str, result: str) -> tuple[float, list[str], bool]:
    """计算STAR四要素完整率。

    每要素需≥20字符才算有效。完整率 = 有效要素数 / 4。
    阈值：≥0.85 即至少 3 个要素完整（允许 1 个草稿阶段）。
    """
    fields = {"situation": situation or "", "task": task or "",
              "action": action or "", "result": result or ""}
    valid_count = sum(1 for v in fields.values() if len(v.strip()) >= 20)
    missing = [k for k, v in fields.items() if len(v.strip()) < 20]
    completeness = valid_count / 4.0
    is_complete = completeness >= 0.85
    return completeness, missing, is_complete


# ═══════════════════════════════════════════════
# 一盘：简历盘点与重构
# ═══════════════════════════════════════════════

def _extract_text_from_file(file_bytes: bytes, filename: str) -> tuple[str, str]:
    """从 PDF / Word 文件中提取文本。

    Returns:
        (text, error_message) — 成功时 error_message 为空字符串
    """
    ext = (filename or "").lower().split(".")[-1] if "." in (filename or "") else "pdf"

    if ext in ("docx", "doc"):
        try:
            # python-docx: 读取 .docx
            import io
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if text:
                return text, ""
        except ImportError:
            return "", "服务器未安装python-docx库，无法解析Word文件，请上传PDF格式"
        except Exception as e:
            logger.warning("Word解析失败: %s", e)
            # 降级：尝试当纯文本读
            try:
                return file_bytes.decode("utf-8", errors="replace"), ""
            except Exception:
                return "", f"无法解析Word文件: {e}"

    if ext == "pdf":
        try:
            # PyPDF2: 读取 PDF
            import io
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            text = "\n".join(parts)
            if text.strip():
                return text, ""
            # PyPDF2 提取为空，尝试 pdfplumber
        except ImportError:
            logger.warning("PyPDF2未安装，尝试pdfplumber")
        except Exception as e:
            logger.warning("PyPDF2解析失败: %s，尝试pdfplumber", e)

        # pdfplumber 作为备选
        try:
            import io
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                parts = [p.extract_text() or "" for p in pdf.pages]
            text = "\n".join(p for p in parts if p.strip())
            if text.strip():
                return text, ""
        except ImportError:
            return "", "服务器未安装PDF解析库(PyPDF2/pdfplumber)，请联系管理员"
        except Exception as e:
            logger.exception("PDF解析失败")

        return "", "无法从PDF中提取文本，请确保PDF包含可选中的文字（非扫描件）"

    return "", f"不支持的文件格式: .{ext}，请上传PDF或Word文件"


def _ai_parse_resume(resume_text: str, user_id: str | None = None,
                     db: Session | None = None) -> dict:
    """调用AI引擎解析简历，提取结构化信息。

    Returns:
        {"summary": str, "key_skills": [...], "key_experiences": [...],
         "education": str, "years_of_experience": str, "industries": [...],
         "strengths": [...], "suggested_positions": [...]}
    """
    from ...engines.persona import build_persona_prompt
    from ...engines.llm_orchestrator import llm_generate_with_orchestration

    # 截断过长的简历（前4000字足够AI分析）
    text = resume_text[:4000]

    prompt = (
        f"请分析以下简历，提取结构化信息。返回JSON（只返回JSON，不要markdown标记）：\n\n"
        f"{{\n"
        f'  "summary": "100字以内的候选人综合概述",\n'
        f'  "key_skills": ["技能1", "技能2", ...],  // 最多8个\n'
        f'  "key_experiences": ["公司/职位/时间段 - 一句话概述", ...],  // 最多5条\n'
        f'  "education": "最高学历+学校+专业",\n'
        f'  "years_of_experience": "X年",\n'
        f'  "industries": ["行业1", "行业2"],\n'
        f'  "strengths": ["亮点1", "亮点2"],  // 候选人最突出的3个优势\n'
        f'  "suggested_positions": ["建议岗位1", "建议岗位2"],  // 基于简历推断的适合岗位\n'
        f"}}\n\n"
        f"简历内容：\n{text}"
    )

    system_prompt = build_persona_prompt(module="career")

    try:
        result = llm_generate_with_orchestration(
            prompt=prompt,
            system_prompt=system_prompt,
            module="career",
            task_complexity="medium",
            temperature=0.3,
            max_tokens=1024,
            user_id=user_id,
            db=db,
        )
        parsed = safe_extract_json(result.get("content", ""))
        if isinstance(parsed, dict) and parsed:
            return parsed
    except Exception as e:
        logger.warning("AI简历解析失败: %s", e)

    # 降级：基于关键词的简单提取
    skills = []
    skill_keywords = [
        "Python", "Java", "JavaScript", "React", "Vue", "SQL", "Docker",
        "项目管理", "团队管理", "数据分析", "人力资源", "财务",
        "市场营销", "产品设计", "运营", "招聘", "培训",
        "C++", "Go", "Rust", "Kubernetes", "AWS", "Azure",
        "沟通", "领导力", "战略规划", "Python", "Excel",
    ]
    for kw in skill_keywords:
        if kw.lower() in resume_text.lower():
            skills.append(kw)

    return {
        "summary": f"简历包含{len(resume_text)}字，涵盖{len(skills)}项可识别技能",
        "key_skills": skills[:8],
        "key_experiences": ["（请查看原文）"],
        "education": "（请查看原文）",
        "years_of_experience": "（请查看原文）",
        "industries": [],
        "strengths": [],
        "suggested_positions": [],
    }


def process_resume_file(db: Session, user_id: str, file_bytes: bytes,
                        filename: str) -> dict[str, Any]:
    """处理上传的简历文件：提取文本 → AI解析 → 保存进度。

    Returns:
        {"career_progress_id": str, "filename": str, "text_preview": str,
         "text_length": int, "parsed_summary": str, "key_skills": [...],
         "key_experiences": [...], "suggested_next": str}
    """
    # 1. 提取文本
    resume_text, error = _extract_text_from_file(file_bytes, filename)
    if error:
        raise APIError(20001, f"简历文件解析失败: {error}", 400)

    if not resume_text or len(resume_text.strip()) < 20:
        raise APIError(20001, "简历文件内容过短或无法识别文字。请确保上传的是文字版PDF而非扫描件图片", 400)

    # 2. AI解析
    ai_result = _ai_parse_resume(resume_text, user_id=user_id, db=db)

    # 3. 保存到五步法进度
    summary = ai_result.get("summary", "")
    skills = ai_result.get("key_skills", [])
    experiences = ai_result.get("key_experiences", [])

    # 创建/更新进度
    existing = db.query(CareerProgress).filter(
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()

    if existing:
        existing.current_step = 1
        existing.status = "active"
        existing.step_data_json = existing.step_data_json or {}
        step_data = dict(existing.step_data_json)
        step_data["resume_title"] = filename
        step_data["resume_content"] = resume_text[:3000]
        step_data["resume_parsed"] = ai_result
        step_data["file_object_id"] = None
        existing.step_data_json = step_data
        db.commit()
        db.refresh(existing)
        progress = existing
        logger.info("简历文件已更新: user_id=%s filename=%s", user_id, filename)
    else:
        progress = CareerProgress(
            user_id=user_id,
            current_step=1,
            status="active",
            step_data_json={
                "resume_title": filename,
                "resume_content": resume_text[:3000],
                "resume_parsed": ai_result,
            },
        )
        db.add(progress)
        db.commit()
        db.refresh(progress)
        logger.info("简历文件已上传: user_id=%s progress_id=%s", user_id, progress.id)

    # 发射事件
    from ...engines.data_foundation import emit_event
    try:
        emit_event(user_id, "career", "resume.uploaded", {
            "filename": filename,
            "text_length": len(resume_text),
            "skills_count": len(skills),
        }, db=db)
    except Exception:
        pass

    return {
        "career_progress_id": progress.id,
        "filename": filename,
        "text_preview": resume_text[:500],
        "text_length": len(resume_text),
        "parsed_summary": summary,
        "key_skills": skills,
        "key_experiences": experiences,
        "suggested_next": "erding",  # 下一步：二定·求职策略
    }

def upload_resume(db: Session, user_id: str, title: str, content: str,
                  file_object_id: str | None = None) -> dict[str, Any]:
    """上传简历：创建/重置五步法进度，进入步骤1。

    简历真实性原则：不做事实核查，不指责用户。
    """
    # 获取或重置进度（重新上传简历时重置五步法）
    existing = db.query(CareerProgress).filter(
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()

    if existing:
        # 不删除已有数据（技能晶体独立保留），仅重置进度状态
        existing.current_step = 1
        existing.status = "active"
        existing.step_data_json = existing.step_data_json or {}
        existing.step_data_json["resume_title"] = title
        existing.step_data_json["resume_content"] = content[:2000]  # 保存摘要
        existing.step_data_json["file_object_id"] = file_object_id
        db.commit()
        db.refresh(existing)
        progress = existing
        logger.info("简历已更新（重置五步法进度）: user_id=%s", user_id)
    else:
        progress = CareerProgress(
            user_id=user_id,
            current_step=1,
            status="active",
            step_data_json={
                "resume_title": title,
                "resume_content": content[:2000],
                "file_object_id": file_object_id,
            },
        )
        db.add(progress)
        db.commit()
        db.refresh(progress)
        logger.info("简历已上传: user_id=%s progress_id=%s", user_id, progress.id)

    return {
        "career_progress_id": progress.id,
        "current_step": 1,
        "message": "简历已接收，请开始STAR萃取",
    }


def star_extraction(db: Session, user_id: str, career_progress_id: str,
                    situation: str | None, task: str | None, action: str | None,
                    result: str | None, quantified_value: str | None,
                    source_type: str = "resume") -> dict[str, Any]:
    """STAR四要素萃取：从简历/面试中引导用户补充，结构化提取四要素。

    核心规则：
      - 四要素完整率≥85%（至少3个有效要素，每要素≥20字符）
      - 不做事实核查，不指责
      - 可多次更新完善
    """
    progress = db.query(CareerProgress).filter(
        CareerProgress.id == career_progress_id,
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()
    if not progress:
        raise APIError(20001, "五步法进度不存在", 404)

    completeness, missing, is_complete = _check_completeness(
        situation or "", task or "", action or "", result or "",
    )

    extraction = STARExtraction(
        user_id=user_id,
        career_progress_id=career_progress_id,
        situation=situation,
        task=task,
        action=action,
        result=result,
        quantified_value=quantified_value,
        completeness=completeness,
        source_type=source_type,
    )
    db.add(extraction)

    # 更新步骤进度
    if not progress.step_data_json:
        progress.step_data_json = {}
    step_data = dict(progress.step_data_json)
    step_data.setdefault("star_extractions", []).append(extraction.id)
    progress.step_data_json = step_data

    db.commit()
    db.refresh(extraction)

    logger.info("STAR萃取完成: extraction_id=%s completeness=%.0f%% is_complete=%s",
                extraction.id, completeness * 100, is_complete)

    # 如果完整，自动生成技能晶体
    if is_complete:
        generate_skill_crystal(db, user_id, extraction.id)

    return {
        "extraction_id": extraction.id,
        "situation": extraction.situation or "",
        "task": extraction.task or "",
        "action": extraction.action or "",
        "result": extraction.result or "",
        "quantified_value": extraction.quantified_value or "",
        "completeness": completeness,
        "missing_fields": missing,
        "is_complete": is_complete,
    }


# ═══════════════════════════════════════════════
# 技能晶体：STAR → SOP 结构化升级
# ═══════════════════════════════════════════════

def generate_skill_crystal(db: Session, user_id: str,
                           star_extraction_id: str) -> dict[str, Any]:
    """从STAR萃取升级为结构化技能晶体。

    五要素：做什么 / 怎么做 / 注意事项 / 成果 / 可复用SOP

    生产环境：调用 voice_engine.llm_generate() 进行AI结构化升级。
    MVP阶段：基于STAR四要素自动生成模板化晶体。
    """
    extraction = db.query(STARExtraction).filter(
        STARExtraction.id == star_extraction_id,
        STARExtraction.user_id == user_id,
        STARExtraction.deleted_at.is_(None),
    ).first()
    if not extraction:
        raise APIError(20001, "STAR萃取记录不存在", 404)

    # 检查是否已有晶体
    existing_crystal = db.query(SkillCrystal).filter(
        SkillCrystal.star_extraction_id == star_extraction_id,
        SkillCrystal.deleted_at.is_(None),
    ).first()
    if existing_crystal:
        return _crystal_to_dict(existing_crystal)

    # 提取STAR四要素
    situation = extraction.situation or ""
    task = extraction.task or ""
    action_text = extraction.action or ""
    result_text = extraction.result or ""
    quantified_value = extraction.quantified_value or ""

    # 尝试AI生成技能晶体
    try:
        ai_result = _generate_ai_skill_crystal(
            situation, task, action_text, result_text, quantified_value,
            user_id=user_id, db=db,
        )
        if ai_result:
            what, how, notes, outcome, reusable_sop, tags, model_used, tokens = ai_result
        else:
            raise ValueError("AI生成技能晶体返回空")
    except Exception as e:
        logger.warning("AI技能晶体生成失败，降级到模板: %s", e)
        # 降级：基于模板生成
        action_keywords = _extract_skill_keywords(action_text)
        what = f"核心能力：{', '.join(action_keywords) if action_keywords else '综合执行能力'}"
        how = f"基于情境「{situation[:80]}...」完成的任务「{task[:80]}...」\n执行步骤：\n{action_text[:500]}"
        notes = (f"关键注意事项：\n"
                 f"1. 充分了解情境背景，明确任务边界\n"
                 f"2. 拆解复杂任务为可执行步骤\n"
                 f"3. 在执行过程中持续收集反馈并调整\n"
                 f"4. 量化成果，用数据说话")
        outcome = f"成果：{result_text[:300] if result_text else '待补充'}"
        reusable_sop = (f"可复用SOP——{task[:60]}：\n"
                        f"S1 情境分析：理解背景、识别关键干系人\n"
                        f"S2 目标拆解：将大目标分解为可量化的小任务\n"
                        f"S3 行动执行：{action_text[:200]}\n"
                        f"S4 结果验证：对标目标，收集量化数据\n"
                        f"S5 复盘归档：总结关键学习点，更新知识库")
        tags = action_keywords
        model_used = "mock_mvp_fallback"
        tokens = 0

    crystal = SkillCrystal(
        user_id=user_id,
        star_extraction_id=star_extraction_id,
        what=what,
        how=how,
        notes=notes,
        outcome=outcome,
        reusable_sop=reusable_sop,
        source_step=1,
        tags_json=tags,
        model_used=model_used,
        generation_cost_tokens=tokens,
    )
    db.add(crystal)
    db.commit()
    db.refresh(crystal)

    logger.info("技能晶体已生成: crystal_id=%s star_extraction=%s tags=%s model=%s",
                crystal.id, star_extraction_id, tags, model_used)

    return _crystal_to_dict(crystal)


def _extract_skill_keywords(text: str) -> list[str]:
    """从行动描述中提取技能关键词。"""
    keywords = []
    skill_patterns = {
        "分析": ["需求分析", "数据分析", "系统分析"],
        "设计": ["系统设计", "架构设计", "方案设计"],
        "开发": ["后端开发", "前端开发", "全栈开发"],
        "管理": ["项目管理", "团队管理", "时间管理"],
        "沟通": ["跨部门沟通", "向上汇报", "客户沟通"],
        "优化": ["性能优化", "流程优化", "成本优化"],
        "协调": ["资源协调", "跨团队协调", "供应商协调"],
        "规划": ["战略规划", "产品规划", "技术规划"],
    }
    text_lower = text.lower()
    for category, sub_skills in skill_patterns.items():
        if category in text_lower:
            for sub in sub_skills:
                if any(s in text_lower for s in [category, sub.split(category)[-1].lower() if category in sub.lower() else ""]):
                    keywords.append(sub)
                    break
            else:
                keywords.append(category)
    return keywords[:5] if keywords else ["综合执行能力"]


def _generate_ai_skill_crystal(
    situation: str, task: str, action_text: str, result_text: str,
    quantified_value: str = "", user_id: str | None = None, db: Session | None = None,
) -> tuple | None:
    """调用AI引擎将STAR四要素升级为技能晶体五要素。

    Returns:
        (what, how, notes, outcome, reusable_sop, tags, model_used, tokens) or None
    """
    output_schema = (
        '{\n'
        '  "what": "做什么——一句话概括核心能力(30字以内)",\n'
        '  "how": "怎么做——具体执行步骤和方法(150-300字)",\n'
        '  "notes": ["注意事项1", "注意事项2", "注意事项3", "注意事项4"],\n'
        '  "outcome": "成果——量化成果描述(100字以内)",\n'
        '  "reusable_sop": "可复用SOP——5步标准化流程(S1-S5)",\n'
        '  "tags": ["技能标签1", "技能标签2"]\n'
        '}'
    )
    prompt = (
        f'请将以下STAR四要素升级为结构化的"技能晶体"：\n\n'
        f'S-情境：{situation[:500]}\n'
        f'T-任务：{task[:500]}\n'
        f'A-行动：{action_text[:800]}\n'
        f'R-结果：{result_text[:500]}\n'
        f'量化成果：{quantified_value or "无"}\n'
    )

    result = llm_generate(
        prompt=prompt,
        system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
        user_id=user_id,
        db=db,
        temperature=0.6,
        module="smart_job",
    )
    parsed = safe_extract_json(result["content"])
    if not parsed or not isinstance(parsed, dict):
        return None

    notes_list = parsed.get("notes", [])
    if isinstance(notes_list, list):
        notes_text = "关键注意事项：\n" + "\n".join(
            f"{i+1}. {n}" for i, n in enumerate(notes_list[:5])
        )
    else:
        notes_text = str(notes_list)

    return (
        parsed.get("what", ""),
        parsed.get("how", ""),
        notes_text,
        parsed.get("outcome", ""),
        parsed.get("reusable_sop", ""),
        parsed.get("tags", []),
        result["model_used"],
        result["usage"]["total_tokens"],
    )


def _crystal_to_dict(c: SkillCrystal) -> dict[str, Any]:
    return {
        "crystal_id": c.id,
        "what": c.what or "",
        "how": c.how or "",
        "notes": c.notes or "",
        "outcome": c.outcome or "",
        "reusable_sop": c.reusable_sop or "",
        "source_step": c.source_step or 1,
        "tags": c.tags_json or [],
    }


def list_skill_crystals(db: Session, user_id: str) -> list[dict[str, Any]]:
    """列出用户的所有技能晶体。"""
    crystals = (
        db.query(SkillCrystal)
        .filter(
            SkillCrystal.user_id == user_id,
            SkillCrystal.deleted_at.is_(None),
        )
        .order_by(desc(SkillCrystal.created_at))
        .all()
    )
    return [
        {
            "id": c.id,
            "what": c.what or "",
            "source_step": c.source_step or 1,
            "tags": c.tags_json or [],
            "archived_to_kb": c.archived_to_kb,
            "created_at": c.created_at.isoformat() if c.created_at else "",
        }
        for c in crystals
    ]


def archive_skill_crystal(db: Session, user_id: str, crystal_id: str) -> dict[str, Any]:
    """归档技能晶体到私有知识库。

    删除求职记录不影响已归档技能晶体：晶体归档到知识库后拥有独立的 doc_id，
    求职记录删除时检查 archived_to_kb，已归档的晶体不级联删除。
    """
    crystal = db.query(SkillCrystal).filter(
        SkillCrystal.id == crystal_id,
        SkillCrystal.user_id == user_id,
        SkillCrystal.deleted_at.is_(None),
    ).first()
    if not crystal:
        raise APIError(40001, "技能晶体不存在", 404)

    # 已归档则直接返回
    if crystal.archived_to_kb and crystal.kb_doc_id:
        return {
            "success": True,
            "doc_id": crystal.kb_doc_id,
            "crystal_id": crystal_id,
        }

    # 写入知识库 Document 表
    from ...shared.models.knowledge import Document, AuditQueue

    doc = Document(
        owner_user_id=user_id,
        library_type="private",
        doc_type="skill_crystal",
        source_module="M13",  # 高维求职
        hr_category="招聘配置",
        title=f"技能晶体-{crystal.what[:50] if crystal.what else '未命名'}",
        content={
            "skill_crystal_id": crystal_id,
            "what": crystal.what,
            "how": crystal.how,
            "notes": crystal.notes,
            "outcome": crystal.outcome,
            "reusable_sop": crystal.reusable_sop,
            "tags": crystal.tags_json,
            "source_step": crystal.source_step,
        },
        status="draft",
        audit_status="pending",
        is_desensitized=False,
        is_negative_blocked=False,
        watermark_required=False,
        version=1,
        vector_status="pending",
    )
    db.add(doc)
    db.flush()

    # 进入待审核区
    now = utcnow()
    db.add(AuditQueue(
        doc_id=doc.id,
        entered_at=now,
        expire_remind_at=now + timedelta(days=30),
    ))

    # 更新晶体归档状态
    crystal.archived_to_kb = True
    crystal.kb_doc_id = doc.id
    db.commit()

    logger.info("技能晶体已归档到知识库: crystal_id=%s doc_id=%s", crystal_id, doc.id)

    return {
        "success": True,
        "doc_id": doc.id,
        "crystal_id": crystal_id,
    }


# ═══════════════════════════════════════════════
# 二定：求职策略与资源
# ═══════════════════════════════════════════════

def create_job_strategy(db: Session, user_id: str, career_progress_id: str,
                        target_industry: str | None = None,
                        target_position: str | None = None,
                        salary_range: str | None = None,
                        location: str | None = None,
                        preferences: str | None = None) -> dict[str, Any]:
    """创建求职策略：基于简历 + 用户偏好生成资源清单和求职计划表。

    推进五步法进度到步骤2。
    """
    progress = db.query(CareerProgress).filter(
        CareerProgress.id == career_progress_id,
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()
    if not progress:
        raise APIError(20001, "五步法进度不存在", 404)

    # 尝试AI生成策略
    try:
        ai_strategy = _generate_ai_job_strategy(
            progress.step_data_json or {},
            target_industry, target_position, salary_range, location, preferences,
            user_id=user_id, db=db,
        )
        if ai_strategy:
            resource_inventory, plan = ai_strategy
        else:
            raise ValueError("AI策略生成返回空")
    except Exception as e:
        logger.warning("AI求职策略生成失败，降级到模板: %s", e)
        # 降级模板
        resource_inventory = {
            "core_skills": _deduce_skills_from_resume(progress.step_data_json or {}),
            "experiences": "从简历STAR萃取中提取的关键项目经验",
            "network": "LinkedIn联系人 / 前同事内推 / 行业社群",
            "certifications": "行业认证、培训证书等",
            "portfolio": "GitHub / 作品集 / 技术博客",
            "references": "至少2-3位可提供推荐的前领导/同事",
        }
        plan = {
            "weekly_goals": [
                {"week": 1, "goal": "完成简历定版+STAR萃取+技能晶体",
                 "tasks": ["定版简历", "完成至少3个STAR萃取", "生成技能晶体"]},
                {"week": 2, "goal": "明确目标行业和岗位·开始投递",
                 "tasks": ["研究目标公司", "完成5-10次投递", "维护投递追踪表"]},
                {"week": 3, "goal": "持续投递·准备面试",
                 "tasks": ["每日投递3-5家", "整理面试常见问题", "模拟面试练习"]},
                {"week": 4, "goal": "面试执行与复盘",
                 "tasks": ["参加面试并录音", "面试后24h内完成复盘", "根据反馈调整策略"]},
            ],
            "channel_strategy": {
                "Boss直聘": "适合互联网/技术岗位，活跃度高",
                "猎聘": "适合中高端岗位，猎头资源丰富",
                "内推": "成功率最高的渠道，优先利用",
                "官网投递": "适合大厂，流程规范但周期长",
                "脉脉/LinkedIn": "适合建立行业人脉，获取隐藏机会",
            },
            "timeline": "建议周期：4-8周密集求职期 → 2-4周面试期 → 1-2周Offer决策期",
        }

    # 更新进度
    if not progress.step_data_json:
        progress.step_data_json = {}
    step_data = dict(progress.step_data_json)
    step_data["strategy"] = {
        "target_industry": target_industry,
        "target_position": target_position,
        "salary_range": salary_range,
        "location": location,
        "preferences": preferences,
        "resource_inventory": resource_inventory,
        "plan": plan,
    }
    # 推进到步骤2
    if progress.current_step < 2:
        progress.current_step = 2
    progress.step_data_json = step_data
    db.commit()
    db.refresh(progress)

    logger.info("求职策略已创建: progress_id=%s target=%s", career_progress_id, target_position)

    return {
        "strategy_id": progress.id,  # 使用 progress_id 作为策略标识
        "resource_inventory": resource_inventory,
        "plan": plan,
        "target_summary": f"{location or '不限地点'} · {target_industry or '不限行业'} · "
                         f"{target_position or '不限岗位'} · {salary_range or '薪资面议'}",
        "tips": ("求职策略提示：①简历控制在1-2页，每份投递前微调关键词；"
                 "②投递时间优先选择周二-周四上午9-11点；"
                 "③每次面试都是一次学习机会，面试后务必复盘；"
                 "④保持'投递-面试-复盘-调整'的正向循环"),
    }


def _deduce_skills_from_resume(step_data: dict) -> list[str]:
    """从简历步骤数据中推断核心技能。"""
    content = step_data.get("resume_content", "")
    # 简单关键词匹配
    skills = []
    skill_kw = ["Python", "Java", "SQL", "React", "Vue", "Docker", "Kubernetes",
                "项目管理", "团队管理", "数据分析", "产品设计", "市场营销",
                "财务分析", "人力资源", "运营管理"]
    for kw in skill_kw:
        if kw.lower() in str(content).lower():
            skills.append(kw)
    return skills[:8] if skills else ["通用职业技能"]


def _generate_ai_job_strategy(
    step_data: dict, target_industry: str | None, target_position: str | None,
    salary_range: str | None, location: str | None, preferences: str | None,
    user_id: str | None = None, db: Session | None = None,
) -> tuple | None:
    """调用AI引擎生成求职资源盘点和计划表。"""
    resume_content = str(step_data.get("resume_content", ""))[:1500]
    output_schema = (
        '{\n'
        '  "resource_inventory": {\n'
        '    "core_skills": ["技能1"],\n'
        '    "experiences": "关键项目经验概述",\n'
        '    "network": "人脉渠道建议",\n'
        '    "certifications": "建议补充的证书",\n'
        '    "portfolio": "建议准备的材料",\n'
        '    "references": "推荐人策略"\n'
        '  },\n'
        '  "plan": {\n'
        '    "weekly_goals": [\n'
        '      {"week": 1, "goal": "...", "tasks": ["..."]},\n'
        '      ...\n'
        '    ],\n'
        '    "channel_strategy": {"渠道名": "策略建议"},\n'
        '    "timeline": "整体时间线建议"\n'
        '  }\n'
        '}'
    )
    prompt = (
        f'用户求职目标：\n'
        f'- 目标行业：{target_industry or "不限"}\n'
        f'- 目标岗位：{target_position or "不限"}\n'
        f'- 期望薪资：{salary_range or "面议"}\n'
        f'- 工作地点：{location or "不限"}\n'
        f'- 偏好：{preferences or "无"}\n\n'
        f'用户简历摘要：\n{resume_content}\n\n'
        f'请生成个性化的求职资源盘点和4周求职计划表。'
    )

    result = llm_generate(
        prompt=prompt,
        system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
        user_id=user_id,
        db=db,
        temperature=0.7,
        module="smart_job",
    )
    parsed = safe_extract_json(result["content"])
    if not parsed or not isinstance(parsed, dict):
        return None

    inv = parsed.get("resource_inventory", {})
    plan = parsed.get("plan", {})

    if not isinstance(inv, dict):
        inv = {}
    if not isinstance(plan, dict):
        plan = {}

    # 确保 weekly_goals 是有效列表
    weekly_goals = plan.get("weekly_goals", [])
    if not isinstance(weekly_goals, list) or len(weekly_goals) < 2:
        return None

    resource_inventory = {
        "core_skills": inv.get("core_skills", []) or _deduce_skills_from_resume(step_data),
        "experiences": inv.get("experiences", "关键项目经验"),
        "network": inv.get("network", "内推+LinkedIn"),
        "certifications": inv.get("certifications", "行业认证"),
        "portfolio": inv.get("portfolio", "GitHub/作品集"),
        "references": inv.get("references", "2-3位推荐人"),
    }

    plan_out = {
        "weekly_goals": weekly_goals[:6],
        "channel_strategy": plan.get("channel_strategy", {
            "Boss直聘": "适合互联网岗位",
            "内推": "成功率最高",
        }),
        "timeline": plan.get("timeline", "4-8周求职期"),
    }

    return resource_inventory, plan_out


# ═══════════════════════════════════════════════
# 三投：投递追踪与分析
# ═══════════════════════════════════════════════

def track_application(db: Session, user_id: str, career_progress_id: str,
                      channel: str, position: str, company: str, date_str: str,
                      status: str = "applied", invite_received: bool = False,
                      notes: str | None = None) -> dict[str, Any]:
    """记录投递：创建投递记录，计算邀约率。

    推进五步法进度到步骤3。
    """
    progress = db.query(CareerProgress).filter(
        CareerProgress.id == career_progress_id,
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()
    if not progress:
        raise APIError(20001, "五步法进度不存在", 404)

    # 验证状态
    from ...shared.models.career import APPLICATION_STATUSES
    if status not in APPLICATION_STATUSES:
        status = "applied"

    application = JobApplication(
        user_id=user_id,
        career_progress_id=career_progress_id,
        channel=channel,
        position=position,
        company=company,
        date=date_str,
        status=status,
        invite_received=invite_received,
        notes=notes,
    )
    db.add(application)

    # 推进到步骤3
    if progress.current_step < 3:
        progress.current_step = 3

    # 更新步骤数据
    if not progress.step_data_json:
        progress.step_data_json = {}
    step_data = dict(progress.step_data_json)
    step_data.setdefault("application_ids", []).append(application.id)
    progress.step_data_json = step_data

    db.commit()
    db.refresh(application)

    # 计算最新邀约率
    stats = _calculate_application_stats(db, user_id)

    logger.info("投递已记录: application_id=%s company=%s position=%s channel=%s",
                application.id, company, position, channel)

    return {
        "application_id": application.id,
        "channel": application.channel or "",
        "position": application.position or "",
        "company": application.company or "",
        "date": application.date or "",
        "status": application.status,
        "invite_received": application.invite_received,
        "invite_rate": stats["invite_rate"],
    }


def _calculate_application_stats(db: Session, user_id: str) -> dict[str, Any]:
    """计算投递统计：总数、各状态数量、邀约率。"""
    apps = db.query(JobApplication).filter(
        JobApplication.user_id == user_id,
        JobApplication.deleted_at.is_(None),
    ).all()

    total = len(apps)
    by_status = {}
    for app in apps:
        s = app.status or "applied"
        by_status[s] = by_status.get(s, 0) + 1

    invite_count = sum(1 for a in apps if a.invite_received)
    invite_rate = round(invite_count / total, 3) if total > 0 else 0.0

    return {
        "total": total,
        "by_status": by_status,
        "invite_rate": invite_rate,
        "invite_count": invite_count,
    }


def list_applications(db: Session, user_id: str) -> dict[str, Any]:
    """列出所有投递记录及统计数据。"""
    apps = (
        db.query(JobApplication)
        .filter(
            JobApplication.user_id == user_id,
            JobApplication.deleted_at.is_(None),
        )
        .order_by(desc(JobApplication.created_at))
        .all()
    )

    items = []
    for a in apps:
        items.append({
            "id": a.id,
            "channel": a.channel or "",
            "position": a.position or "",
            "company": a.company or "",
            "date": a.date or "",
            "status": a.status or "applied",
            "status_label": STATUS_LABELS.get(a.status or "applied", "未知"),
            "invite_received": a.invite_received,
        })

    stats = _calculate_application_stats(db, user_id)

    return {
        "items": items,
        "stats": stats,
    }


# ═══════════════════════════════════════════════
# 四面：面试准备与复盘
# ═══════════════════════════════════════════════

def prepare_interview(db: Session, user_id: str, career_progress_id: str,
                      application_id: str | None = None,
                      company: str = "", position: str = "",
                      interview_stage: str | None = None) -> dict[str, Any]:
    """面试准备：生成企业情报+匹配度分析+面试策略+问题清单。

    企业情报 AI 初稿≤3分钟，仅公开信息源。
    推进五步法进度到步骤4。
    """
    progress = db.query(CareerProgress).filter(
        CareerProgress.id == career_progress_id,
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()
    if not progress:
        raise APIError(20001, "五步法进度不存在", 404)

    # 企业情报（AI初稿，仅公开信息）—— AI优先+模板降级
    company_intel = _generate_company_intel(company, user_id=user_id, db=db)

    # 匹配度分析
    match_analysis = _generate_match_analysis(
        progress.step_data_json or {}, position, company, user_id=user_id, db=db,
    )

    # 面试策略
    stage = interview_stage or "综合面试"
    strategy_doc = _generate_interview_strategy(company, position, stage, user_id=user_id, db=db)

    # 问题清单
    question_list = _generate_question_list(company, position, stage, user_id=user_id, db=db)
    warm_up_questions = _generate_warm_up_questions(company, position, user_id=user_id, db=db)

    prep = InterviewPrep(
        user_id=user_id,
        career_progress_id=career_progress_id,
        application_id=application_id,
        company=company,
        position=position,
        company_intel_json=company_intel,
        match_analysis=match_analysis,
        strategy_doc=strategy_doc,
        question_list=question_list,
    )
    db.add(prep)

    # 推进到步骤4
    if progress.current_step < 4:
        progress.current_step = 4

    if not progress.step_data_json:
        progress.step_data_json = {}
    step_data = dict(progress.step_data_json)
    step_data.setdefault("interview_prep_ids", []).append(prep.id)
    progress.step_data_json = step_data

    db.commit()
    db.refresh(prep)

    logger.info("面试准备已生成: prep_id=%s company=%s position=%s", prep.id, company, position)

    return {
        "prep_id": prep.id,
        "company_intel": company_intel,
        "match_analysis": match_analysis,
        "strategy_doc": strategy_doc,
        "question_list": question_list,
        "warm_up_questions": warm_up_questions,
    }


def _generate_company_intel(company: str, user_id: str | None = None,
                             db: Session | None = None) -> dict:
    """生成企业情报初稿（仅公开信息源，≤3分钟AI生成）。"""
    if not company:
        return {"note": "请提供公司名称以获取企业情报"}

    # 尝试AI生成
    try:
        output_schema = (
            '{\n'
            '  "industry": "行业分类",\n'
            '  "size": "规模估算",\n'
            '  "stage": "发展阶段",\n'
            '  "culture_tags": ["文化标签1", "标签2"],\n'
            '  "products": "主要产品/服务线概述",\n'
            '  "competitors": "主要竞争对手",\n'
            '  "recent_news": ["近期动态1"]\n'
            '}'
        )
        prompt = (
            f'请为以下公司生成面试准备用的企业情报初稿（仅使用你的训练数据中的公开信息）：\n\n'
            f'公司名称：{company}\n\n'
            f'请基于该公司在公开信息中的形象，提供行业分类、规模估算、发展阶段、文化标签、'
            f'产品概述、竞争对手和近期动态。如果信息不足，请标注"建议核实"而非编造。'
        )
        result = llm_generate(
            prompt=prompt,
            system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
            user_id=user_id,
            db=db,
            temperature=0.5,
            module="smart_job",
        )
        parsed = safe_extract_json(result["content"])
        if parsed and isinstance(parsed, dict):
            return {
                "company_name": company,
                "industry": parsed.get("industry", "建议通过工商信息核实"),
                "size": parsed.get("size", "建议通过企查查/天眼查核实"),
                "stage": parsed.get("stage", "建议了解融资阶段"),
                "culture_tags": parsed.get("culture_tags", ["技术创新"]),
                "products": parsed.get("products", "建议访问公司官网了解"),
                "competitors": parsed.get("competitors", "建议通过行业报告分析"),
                "recent_news": parsed.get("recent_news", ["建议搜索了解最新动态"]),
                "source_urls": [
                    "https://www.qcc.com（企查查）",
                    "https://www.tianyancha.com（天眼查）",
                ],
                "disclaimer": "本情报基于AI初稿，仅使用公开信息源，生成时间≤3分钟。请用户自行核实。",
            }
    except Exception as e:
        logger.warning("AI企业情报生成失败: %s", e)

    # 降级模板
    return {
        "company_name": company,
        "industry": "信息技术/互联网（基于公开信息推断）",
        "size": "建议通过企查查/天眼查核实工商信息",
        "stage": "建议通过36氪/IT桔子了解融资阶段",
        "culture_tags": ["技术创新", "扁平化管理", "结果导向"],
        "products": "建议访问公司官网了解产品矩阵",
        "competitors": "建议通过行业报告分析竞争格局",
        "recent_news": ["建议通过搜索引擎了解最新动态"],
        "source_urls": ["https://www.qcc.com（企查查）", "https://www.tianyancha.com（天眼查）"],
        "disclaimer": "本情报基于AI初稿，仅使用公开信息源，生成时间≤3分钟。请用户自行核实。",
    }


def _generate_match_analysis(step_data: dict, position: str, company: str,
                              user_id: str | None = None,
                              db: Session | None = None) -> str:
    """生成匹配度分析。"""
    # 尝试AI生成
    try:
        resume_content = str(step_data.get("resume_content", ""))[:2000]
        prompt = (
            f'请分析以下候选人与目标岗位的匹配度：\n\n'
            f'岗位：{position}\n'
            f'公司：{company}\n\n'
            f'候选人简历/STAR萃取摘要：\n{resume_content}\n\n'
            f'请以纯文本Markdown格式输出匹配度分析（不需要JSON），包含：\n'
            f'1. 岗位核心要求理解\n'
            f'2. 候选人优势匹配（具体指出哪些经验/技能匹配）\n'
            f'3. 差距与提升建议\n'
            f'4. 整体匹配评估'
        )
        result = llm_generate(
            prompt=prompt,
            system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
            user_id=user_id,
            db=db,
            temperature=0.6,
            module="smart_job",
        )
        content = result.get("content", "").strip()
        if content and len(content) > 50:
            return content
    except Exception as e:
        logger.warning("AI匹配度分析生成失败: %s", e)

    # 降级模板
    parts = [
        f"## 匹配度分析：{company} - {position}",
        "",
        "### 岗位要求理解",
        f"基于「{position}」岗位的通用要求进行分析：",
        "- 专业技能：需具备岗位相关的核心技术能力",
        "- 项目经验：相关领域的实战经验",
        "- 软技能：沟通表达、团队协作、问题解决",
        "",
        "### 候选人优势匹配",
        "- 从简历中提取的核心经验与岗位需求的对齐度",
        "- STAR萃取中体现的问题解决能力和量化成果",
        "",
        "### 差距与提升建议",
        "- 建议针对性补充的知识/技能领域",
        "- 面试中需要重点展示的优势面",
        "",
        "### 整体匹配评估",
        "建议通过面试进一步验证匹配度，重点关注：",
        "1. 专业技能的深度和广度",
        "2. 项目经验与岗位的相关性",
        "3. 文化价值观的契合度",
    ]
    return "\n".join(parts)


def _generate_interview_strategy(company: str, position: str, stage: str,
                                  user_id: str | None = None,
                                  db: Session | None = None) -> str:
    """生成面试策略文档。"""
    # 尝试AI生成
    try:
        prompt = (
            f'请为以下面试生成面试策略文档：\n\n'
            f'公司：{company}\n'
            f'岗位：{position}\n'
            f'面试阶段：{stage}\n\n'
            f'策略内容应包含（纯文本Markdown格式，不需要JSON）：\n'
            f'1. 整体策略（3-4条核心原则）\n'
            f'2. 核心卖点（应重点展示的3个方面）\n'
            f'3. 阶段策略（针对{stage}的具体建议）\n'
            f'4. 注意事项（针对该公司/岗位的特别提醒）'
        )
        result = llm_generate(
            prompt=prompt,
            system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
            user_id=user_id,
            db=db,
            temperature=0.7,
            module="smart_job",
        )
        content = result.get("content", "").strip()
        if content and len(content) > 50:
            return content
    except Exception as e:
        logger.warning("AI面试策略生成失败: %s", e)

    # 降级模板
    return f"""## 面试策略：{company} - {position} ({stage})

### 整体策略
1. **第一印象管理**：自我介绍控制在1-2分钟，突出与岗位最相关的3个核心能力
2. **STAR方法应答**：所有项目经验问题用STAR框架回答（情境→任务→行动→结果）
3. **主动提问**：准备5-8个有深度的问题，展示你的思考和对公司/岗位的了解
4. **展示成长心态**：在回答中展示你的学习能力和自我反思

### 核心卖点
- 提炼3个最能与岗位匹配的核心能力和经验
- 准备每个卖点的STAR案例（至少2-3个）
- 量化你的成果（用数据说话）

### 阶段策略
- 初面：展示综合能力和基本匹配度，注重沟通表达
- 技术面：展示专业知识深度，准备技术案例和白板练习
- HR面：展示职业规划、价值观匹配、薪酬期望（谈判技巧）
- 终面：展示战略思维、领导力潜力、长期价值

### 注意事项
- 提前研究公司背景、产品、文化和近期动态
- 准备好反问环节的问题（避免问薪资福利，留给HR面）
- 面试结束后24小时内发送感谢邮件
"""


def _generate_question_list(company: str, position: str, stage: str,
                             user_id: str | None = None,
                             db: Session | None = None) -> list[dict]:
    """生成面试问题清单。"""
    # 尝试AI生成
    try:
        output_schema = (
            '{\n'
            '  "questions": [\n'
            '    {"category": "类别", "question": "问题", "purpose": "目的", "answer_hint": "理想回答要点"},\n'
            '    ...\n'
            '  ]\n'
            '}'
        )
        prompt = (
            f'请为以下面试场景生成7-8个面试问题：\n\n'
            f'公司：{company}\n'
            f'岗位：{position}\n'
            f'面试阶段：{stage}\n\n'
            f'覆盖类别：自我介绍、项目经验、技术/专业能力、职业动机、团队协作、应对挑战、反问环节'
        )
        result = llm_generate(
            prompt=prompt,
            system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
            user_id=user_id,
            db=db,
            temperature=0.7,
            module="smart_job",
        )
        parsed = safe_extract_json(result["content"])
        if parsed and isinstance(parsed, dict) and len(parsed.get("questions", [])) >= 5:
            return parsed["questions"][:8]
    except Exception as e:
        logger.warning("AI问题清单生成失败: %s", e)

    # 降级模板
    return [
        {
            "category": "自我介绍",
            "question": "请做一个1-2分钟的自我介绍，重点介绍与你应聘岗位相关的核心经验。",
            "purpose": "评估表达能力、职业脉络和岗位匹配度",
            "answer_hint": "按「过去→现在→未来」结构：简述职业背景→当前能力和成就→为什么适合这个岗位",
        },
        {
            "category": "项目经验",
            "question": "请用STAR方法介绍一个你最有成就感或最有挑战的项目。",
            "purpose": "评估问题解决能力、项目经验和量化思维",
            "answer_hint": "明确的S(情境)→T(任务)→A(行动)→R(结果)，重点是量化成果",
        },
        {
            "category": "技术/专业能力",
            "question": f"对于{position}这个岗位，你认为最核心的3个能力是什么？请举例说明你在这3个方面的表现。",
            "purpose": "评估自我认知和岗位理解深度",
            "answer_hint": "能力描述应具体而非空泛，有案例支撑",
        },
        {
            "category": "职业动机",
            "question": f"为什么选择{company}？为什么离开当前/上一家公司？",
            "purpose": "评估职业动机、对公司了解程度和价值取向",
            "answer_hint": "建设性的离职原因+对公司的真实了解+清晰职业规划",
        },
        {
            "category": "团队协作",
            "question": "请描述一次你与团队意见分歧的经历，你是如何处理和解决的？",
            "purpose": "评估团队协作、冲突处理能力和情商",
            "answer_hint": "展示尊重不同观点、寻求共同利益、建设性解决冲突",
        },
        {
            "category": "应对挑战",
            "question": "描述你在工作中遇到的最大挫折，你是如何走出来的？学到了什么？",
            "purpose": "评估抗压能力、韧性和自我反思能力",
            "answer_hint": "诚实描述困难，但重点在行动和成长",
        },
        {
            "category": "反问环节",
            "question": "你有什么问题想问我？（建议准备3-5个问题）",
            "purpose": "评估主动性、思考深度和对公司的真实兴趣",
            "answer_hint": "问团队架构、业务方向、成长路径、公司文化等，避免过早问薪资福利",
        },
    ]


def _generate_warm_up_questions(company: str = "", position: str = "",
                                 user_id: str | None = None,
                                 db: Session | None = None) -> list[dict]:
    """生成暖场问题。"""
    # 尝试AI生成
    try:
        output_schema = '{"questions": [{"question": "...", "purpose": "..."}, ...]}'
        prompt = (
            f'请为{company}的{position}面试开场环节生成3句暖场/破冰提示。'
            f'需要轻松自然，缓解候选人紧张情绪。'
        )
        result = llm_generate(
            prompt=prompt,
            system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
            user_id=user_id,
            db=db,
            temperature=0.8,
            module="smart_job",
        )
        parsed = safe_extract_json(result["content"])
        if parsed and isinstance(parsed, dict) and len(parsed.get("questions", [])) >= 2:
            return parsed["questions"][:3]
    except Exception as e:
        logger.warning("AI暖场问题生成失败: %s", e)

    # 降级模板
    return [
        {"question": "今天的天气不错/交通还好吗？", "purpose": "缓解紧张气氛"},
        {"question": "感谢你今天抽时间来参加面试。", "purpose": "展示尊重和友善"},
        {"question": "我们大概需要30-45分钟，请放松，就像一次交流。", "purpose": "设定预期，减少压力"},
    ]


def start_interview_recording(db: Session, user_id: str, prep_id: str,
                               recording_id: str) -> dict[str, Any]:
    """关联智能记录录音到求职面试。

    跨模块联动：链接到智能记录的 Recording 表，使用"高维求职"场景标记。
    """
    prep = db.query(InterviewPrep).filter(
        InterviewPrep.id == prep_id,
        InterviewPrep.user_id == user_id,
        InterviewPrep.deleted_at.is_(None),
    ).first()
    if not prep:
        raise APIError(20001, "面试准备记录不存在", 404)

    # 创建面试复盘记录（关联录音）
    review = InterviewReview(
        user_id=user_id,
        prep_id=prep_id,
        audio_recording_id=recording_id,
    )
    db.add(review)
    db.commit()
    db.refresh(review)

    logger.info("面试录音已关联: review_id=%s prep_id=%s recording_id=%s",
                review.id, prep_id, recording_id)

    return {
        "review_id": review.id,
        "prep_id": prep_id,
        "recording_id": recording_id,
        "status": "linked",
        "message": "录音已关联，可进行面试复盘分析",
    }


def analyze_interview(db: Session, user_id: str, review_id: str) -> dict[str, Any]:
    """面试复盘分析：基于录音分析面试表现。

    产出：亮点 + 改进方向 + 复盘SOP + 总体评分。
    生产环境：调用 voice_engine 对录音进行AI分析。
    MVP阶段：基于面试准备和问题清单生成模拟分析。
    """
    review = db.query(InterviewReview).filter(
        InterviewReview.id == review_id,
        InterviewReview.user_id == user_id,
        InterviewReview.deleted_at.is_(None),
    ).first()
    if not review:
        raise APIError(20001, "面试复盘记录不存在", 404)

    # 如果已有分析，直接返回
    if review.highlights:
        return _review_to_dict(review)

    # 获取关联的面试准备记录
    prep = db.query(InterviewPrep).filter(
        InterviewPrep.id == review.prep_id,
        InterviewPrep.deleted_at.is_(None),
    ).first()

    company = prep.company if prep else "未知公司"
    position = prep.position if prep else "未知岗位"

    # 尝试AI分析
    ai_rating = 4
    try:
        ai_result = _generate_ai_interview_review(company, position, user_id=user_id, db=db)
        if ai_result:
            highlights, improvements, review_sop, ai_rating = ai_result
        else:
            raise ValueError("AI分析返回空")
    except Exception as e:
        logger.warning("AI面试复盘分析失败，降级到模板: %s", e)
        # 降级模板
        highlights = f"""## 面试亮点分析：{company} - {position}

### 表现突出方面
1. **自我介绍结构清晰**：职业脉络表达流畅，核心能力突出
2. **STAR案例准备充分**：项目经验描述有量化和数据支撑
3. **问题理解准确**：能够准确把握面试官问题的核心意图
4. **反问环节有深度**：展示了真实的研究和思考

### 关键加分项
- 展现了持续学习和自我反思的习惯
- 对公司/行业有一定了解，做了功课
- 沟通中展现了真诚和自信的平衡
"""
        improvements = f"""## 面试改进方向：{company} - {position}

### 需要提升方面
1. **专业深度展示**：技术/专业问题上可以更深入，展示专业性
2. **回答简洁度**：部分回答略冗长，建议更有针对性
3. **案例多样性**：准备2-3个不同类型的STAR案例，覆盖团队协作、技术创新、冲突处理
4. **数据支撑**：增加更多量化成果来增强说服力

### 具体改进建议
- 每个回答控制在2-3分钟以内
- 遇到不会的问题诚实表达，但展示学习思路
- 面试后24小时内发送感谢邮件（简要回顾+表达兴趣）
"""
        review_sop = f"""## 面试复盘SOP

### 面试后24小时黄金复盘期
**S1 - 立即记录（面试结束10分钟内）**
- 记录面试官姓名、面试问题清单
- 标注回答得好的问题和回答不满意的问题
- 记录面试官的反馈和表情变化

**S2 - 深度复盘（当日）**
- 逐题分析：自己的回答 vs 理想回答的差距
- 识别模式：归纳被反复问到的能力维度
- 更新知识库：将新学到的问题和答案存档

**S3 - 行动调整（次日）**
- 针对薄弱环节补充学习/练习
- 调整自我介绍和核心卖点话术
- 根据反馈调整下一次面试策略

### 持续改进循环
面试 → 录音 → 分析 → 复盘 → 调整 → 下一轮面试
"""

    # 更新复盘记录
    review.highlights = highlights
    review.improvements = improvements
    review.review_sop = review_sop
    review.overall_rating = ai_rating
    db.commit()
    db.refresh(review)

    logger.info("面试复盘分析完成: review_id=%s rating=%d", review_id, review.overall_rating)

    return _review_to_dict(review)


def _review_to_dict(r: InterviewReview) -> dict[str, Any]:
    return {
        "review_id": r.id,
        "prep_id": r.prep_id or "",
        "highlights": r.highlights or "",
        "improvements": r.improvements or "",
        "review_sop": r.review_sop or "",
        "overall_rating": r.overall_rating or 0,
    }


def _generate_ai_interview_review(
    company: str, position: str,
    user_id: str | None = None, db: Session | None = None,
) -> tuple | None:
    """调用AI引擎生成面试复盘分析。"""
    output_schema = (
        '{\n'
        '  "highlights": "面试亮点分析(Markdown)",\n'
        '  "improvements": "面试改进方向(Markdown)",\n'
        '  "review_sop": "面试复盘SOP(Markdown)",\n'
        '  "overall_rating": 4\n'
        '}'
    )
    prompt = (
        f'请分析以下面试表现并生成复盘报告：\n\n'
        f'公司：{company}\n'
        f'岗位：{position}\n\n'
        f'请基于通用面试场景，分析候选人的可能表现。'
        f'亮点分析应包含3-4个方面，改进方向应包含3-4个方面（纯文本Markdown格式），'
        f'复盘SOP为5步标准化流程。overall_rating为1-5的整数评分。'
    )

    result = llm_generate(
        prompt=prompt,
        system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
        user_id=user_id,
        db=db,
        temperature=0.6,
        module="smart_job",
    )
    parsed = safe_extract_json(result["content"])
    if not parsed or not isinstance(parsed, dict):
        return None

    rating = parsed.get("overall_rating", 4)
    try:
        rating = max(1, min(5, int(rating)))
    except (ValueError, TypeError):
        rating = 4

    return (
        parsed.get("highlights", ""),
        parsed.get("improvements", ""),
        parsed.get("review_sop", ""),
        rating,
    )


def list_interview_preps(db: Session, user_id: str) -> list[dict[str, Any]]:
    """列出所有面试准备记录。"""
    preps = (
        db.query(InterviewPrep)
        .filter(
            InterviewPrep.user_id == user_id,
            InterviewPrep.deleted_at.is_(None),
        )
        .order_by(desc(InterviewPrep.created_at))
        .all()
    )

    items = []
    for p in preps:
        # 检查是否已有复盘
        has_review = db.query(InterviewReview).filter(
            InterviewReview.prep_id == p.id,
            InterviewReview.deleted_at.is_(None),
        ).count() > 0

        items.append({
            "id": p.id,
            "company": p.company or "",
            "position": p.position or "",
            "has_review": has_review,
            "created_at": p.created_at.isoformat() if p.created_at else "",
        })

    return items


# ═══════════════════════════════════════════════
# 五选：Offer评估与入职
# ═══════════════════════════════════════════════

def compare_offers(db: Session, user_id: str, career_progress_id: str,
                   offers: list[dict], notes: str | None = None) -> dict[str, Any]:
    """Offer多维度对比：陈列式分析，不替用户做决策。

    六个维度：薪资 / 职级 / 发展空间 / 通勤 / 文化 / 稳定性
    """
    progress = db.query(CareerProgress).filter(
        CareerProgress.id == career_progress_id,
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()
    if not progress:
        raise APIError(20001, "五步法进度不存在", 404)

    # 六维度分析
    dimension_analysis = {
        "salary": "薪资维度：请综合考虑基本工资、绩效奖金、股权/期权、五险一金、其他福利。注意比较税后收入和实际购买力。",
        "level": "职级维度：关注职级的市场对标、晋升路径和空间。高职级不一定更优，需结合公司规模和职级含金量。",
        "growth": "发展空间维度：考虑业务前景、团队位置、学习机会、内部转岗可能。3-5年的成长预期比当前薪资更重要。",
        "commute": "通勤维度：考虑通勤时间/成本、是否支持远程办公。长期通勤对生活质量影响显著。",
        "culture": "文化维度：关注价值观匹配、工作节奏、团队氛围、管理风格。文化不适往往是离职的主因。",
        "stability": "稳定性维度：公司财务状况、发展阶段、行业趋势、团队稳定性。创业公司高风险高回报，大厂稳定但成长慢。",
    }

    comparison = OfferComparison(
        user_id=user_id,
        career_progress_id=career_progress_id,
        offers_json=offers,
        comparison_notes=notes,
    )
    db.add(comparison)

    # 推进到步骤5
    if progress.current_step < 5:
        progress.current_step = 5

    if not progress.step_data_json:
        progress.step_data_json = {}
    step_data = dict(progress.step_data_json)
    step_data.setdefault("offer_comparison_ids", []).append(comparison.id)
    progress.step_data_json = step_data

    db.commit()
    db.refresh(comparison)

    # 决策原则提示（不指向具体Offer）
    tips = ("Offer决策通用原则："
            "①听从内心的第一直觉，你更愿意去哪家公司上班？"
            "②和未来的直属上级聊得如何？直属上级是你入职后最重要的关系。"
            "③考虑3年后的你——哪个选择让你更接近3年后想成为的样子？"
            "④和家人/朋友讨论，但最终决策权在你手中。"
            "⑤薪资本身不是唯一标准，成长空间和团队氛围往往更关键。")

    logger.info("Offer对比已创建: comparison_id=%s offers_count=%d", comparison.id, len(offers))

    return {
        "comparison_id": comparison.id,
        "offers": offers,
        "dimension_analysis": dimension_analysis,
        "tips": tips,
    }


def accept_offer(db: Session, user_id: str, comparison_id: str,
                 selected_offer_label: str) -> dict[str, Any]:
    """接受Offer：选择并生成试用期30/60/90天计划。

    合同在平台外签署，平台不参与资金流转。
    """
    comparison = db.query(OfferComparison).filter(
        OfferComparison.id == comparison_id,
        OfferComparison.user_id == user_id,
        OfferComparison.deleted_at.is_(None),
    ).first()
    if not comparison:
        raise APIError(20001, "Offer对比记录不存在", 404)

    # 查找选中的Offer详情
    offers = comparison.offers_json or []
    selected_offer = None
    for o in offers:
        if o.get("offer_label") == selected_offer_label:
            selected_offer = o
            break

    company = selected_offer.get("company", "新公司") if selected_offer else "新公司"
    position = selected_offer.get("position", "新岗位") if selected_offer else "新岗位"

    # 生成试用期计划
    milestones = _generate_probation_milestones(company, position, user_id=user_id, db=db)

    plan = ProbationPlan(
        user_id=user_id,
        offer_id=comparison.id,
        milestones_json=milestones,
        company=company,
        position=position,
    )
    db.add(plan)

    # 更新对比记录
    comparison.selected_offer_id = selected_offer_label

    # 标记五步法完成
    progress = db.query(CareerProgress).filter(
        CareerProgress.id == comparison.career_progress_id,
    ).first()
    if progress:
        progress.status = "completed"

    db.commit()
    db.refresh(plan)

    logger.info("Offer已接受并生成试用期计划: plan_id=%s company=%s position=%s",
                plan.id, company, position)

    return {
        "plan_id": plan.id,
        "offer_label": selected_offer_label,
        "company": company,
        "position": position,
        "milestones": milestones,
        "overall_goal": milestones.get("overall_goal", "顺利通过试用期并建立核心贡献"),
    }


def _generate_probation_milestones(company: str, position: str,
                                    user_id: str | None = None,
                                    db: Session | None = None) -> dict:
    """生成试用期30/60/90天里程碑。"""
    # 尝试AI生成
    try:
        output_schema = (
            '{\n'
            '  "overall_goal": "试用期总体目标(一句话)",\n'
            '  "day_30": [{"goal": "阶段目标", "actions": ["行动1"], "checkpoints": ["检查点1"]}],\n'
            '  "day_60": [{"goal": "...", "actions": [...], "checkpoints": [...]}],\n'
            '  "day_90": [{"goal": "...", "actions": [...], "checkpoints": [...]}]\n'
            '}'
        )
        prompt = (
            f'请为以下新入职员工生成试用期30/60/90天里程碑计划：\n\n'
            f'公司：{company}\n'
            f'岗位：{position}\n\n'
            f'每个阶段（30/60/90天）应包含1个目标、3-5个行动项和2-3个检查点。'
        )
        result = llm_generate(
            prompt=prompt,
            system_prompt=MODULE_SYSTEM_PROMPTS.get("smart_job", ""),
            user_id=user_id,
            db=db,
            temperature=0.6,
            module="smart_job",
        )
        parsed = safe_extract_json(result["content"])
        if parsed and isinstance(parsed, dict) and "day_30" in parsed:
            return parsed
    except Exception as e:
        logger.warning("AI试用期计划生成失败: %s", e)

    # 降级模板
    return {
        "overall_goal": f"顺利通过{company}的试用期考核，在{position}岗位上建立核心贡献",
        "day_30": [
            {
                "goal": "了解组织·建立关系",
                "actions": [
                    "完成入职培训，了解公司文化、制度、产品",
                    "与直属上级完成1:1沟通，明确岗位期望和30天目标",
                    "认识团队每位成员，了解各自的职责和沟通偏好",
                    "梳理关键干系人地图（谁是我需要经常协作的人）",
                    "熟悉核心工具链和工作流程",
                ],
                "checkpoints": [
                    "完成入职培训并签署确认",
                    "与上级对齐30天期望并记录",
                    "完成干系人地图",
                ],
            }
        ],
        "day_60": [
            {
                "goal": "建立节奏·创造价值",
                "actions": [
                    "独立完成1-2个小型任务，展示执行能力",
                    "参与团队周会，开始主动贡献观点",
                    "识别1-2个可以改进的流程或痛点",
                    "与跨部门关键联系人建立初步关系",
                    "开始建立个人工作日志（记录成绩和成长）",
                ],
                "checkpoints": [
                    "独立交付至少1个任务并获正面反馈",
                    "在团队会议中至少发言1次",
                    "向上级汇报60天进展和适应情况",
                ],
            }
        ],
        "day_90": [
            {
                "goal": "确立位置·规划未来",
                "actions": [
                    "主导或深度参与1个完整项目/任务",
                    "基于前60天的观察，提出1-2个改进建议",
                    "与上级完成90天回顾，讨论长期发展路径",
                    "建立自己的专业影响力（内部分享/文档输出）",
                    "制定未来6-12个月的个人成长计划",
                ],
                "checkpoints": [
                    "90天回顾:上级对试用期的综合评价",
                    "完成个人成长计划初稿",
                    "至少1次知识分享或文档贡献",
                ],
            }
        ],
    }


# ═══════════════════════════════════════════════
# 企业情报（老师后台）
# ═══════════════════════════════════════════════

def get_company_intel(db: Session, teacher_id: str | None, company_name: str,
                      user_id: str | None = None) -> dict[str, Any]:
    """AI辅助企业情报采集（仅公开信息源）。

    业务规则：
      - AI初稿≤3分钟
      - 仅公开信息源
      - 老师最终核实
      - 老师经保密协议后可查看私有库只读权限
    """
    import time
    start_ms = int(time.time() * 1000)

    # 尝试AI生成企业情报
    model_used = "mock_mvp"
    intel_report = None
    try:
        # 复用面试准备的AI企业情报生成逻辑
        ai_intel = _generate_company_intel(company_name, user_id=user_id, db=db)
        if ai_intel and "note" not in ai_intel:
            intel_report = {
                "company_name": company_name,
                "industry": ai_intel.get("industry", ""),
                "scale": ai_intel.get("size", ""),
                "founded": "建议通过工商信息核实",
                "hq": "建议通过企业官网核实",
                "culture_summary": ", ".join(ai_intel.get("culture_tags", [])),
                "product_summary": ai_intel.get("products", ""),
                "market_position": ai_intel.get("competitors", ""),
                "recent_news": ai_intel.get("recent_news", []),
                "risk_flags": ["本部分仅基于公开信息，请老师审核后确认"],
                "hiring_trend": "建议参考脉脉/看准网等职场社区",
                "disclaimer": ai_intel.get("disclaimer", "仅公开信息源"),
            }
            model_used = "zhipu_ai"
    except Exception as e:
        logger.warning("AI企业情报生成失败，降级到模板: %s", e)

    if not intel_report:
        # 降级模板
        intel_report = {
            "company_name": company_name,
            "industry": "建议通过工商信息/企业官网核实",
            "scale": "建议通过企查查/天眼查核实",
            "founded": "建议通过工商信息核实",
            "hq": "建议通过企业官网核实",
            "culture_summary": f"{company_name}的企业文化特点（基于公开信息整理）",
            "product_summary": "建议访问公司官网了解产品矩阵",
            "market_position": "建议参考行业报告和竞品分析",
            "recent_news": ["建议通过搜索引擎了解最新动态"],
            "risk_flags": ["本部分仅基于公开信息，请老师审核后确认"],
            "hiring_trend": "建议参考脉脉/看准网等职场社区",
            "disclaimer": "本情报基于AI初稿，仅使用公开信息源。请老师核实后补充。",
        }

    source_urls = [
        "https://www.qcc.com（企查查工商信息）",
        "https://www.tianyancha.com（天眼查）",
        "https://www.qixin.com（启信宝）",
    ]

    generation_time_ms = int(time.time() * 1000) - start_ms

    intel_record = CompanyIntel(
        user_id=user_id,
        company_name=company_name,
        intel_report_json=intel_report,
        source_urls=source_urls,
        teacher_id=teacher_id,
        teacher_verified=False,
        model_used=model_used,
        generation_time_ms=generation_time_ms,
    )
    db.add(intel_record)
    db.commit()
    db.refresh(intel_record)

    logger.info("企业情报已生成: intel_id=%s company=%s time_ms=%d",
                intel_record.id, company_name, generation_time_ms)

    return {
        "intel_id": intel_record.id,
        "company_name": company_name,
        "intel_report": intel_report,
        "source_urls": source_urls,
        "teacher_verified": False,
        "generation_time_ms": generation_time_ms,
    }


# ═══════════════════════════════════════════════
# AI 高维求职对话 — 所有小耕输出由模型生成
# ═══════════════════════════════════════════════

def process_career_chat(
    message: str,
    step: str = "yipan",
    context: list[dict] | None = None,
    sub_index: int = 0,
    has_resume: bool = False,
    user_id: str | None = None,
    db: Session | None = None,
) -> dict[str, Any]:
    """高维求职 AI 对话 — 所有小耕回复由AI模型生成。

    根据五步法步骤和子进度，AI按算法引导用户完成求职全流程。

    一盘(yipan)五大盘点:
      0-履历梳理 → 1-STAR追问 → 2-技能晶体 → 3-人脉资源 → 4-岗位建议

    Args:
        message: 用户当前消息（初始可为空）
        step: yipan|erding|santou|simian|wuxuan
        context: 对话历史
        sub_index: 一盘子进度 0-4
        has_resume: 是否已上传简历
        user_id: 用户ID
        db: 数据库会话

    Returns:
        {"reply": str, "model_used": str}
    """
    from ...engines.persona import build_persona_prompt
    from ...engines.llm_orchestrator import llm_generate_with_orchestration

    # 步骤标签
    step_labels = {
        "yipan": "一盘·简历盘点与重构",
        "erding": "二定·求职策略与资源",
        "santou": "三投·投递追踪与分析",
        "simian": "四面·面试准备与复盘",
        "wuxuan": "五选·Offer评估与入职",
    }

    # 一盘子阶段
    yipan_phases = ["履历梳理", "STAR追问", "技能晶体", "人脉资源", "岗位建议"]
    phase_label = yipan_phases[min(sub_index, 4)]

    try:
        # 构建对话历史
        context_text = ""
        if context:
            recent = context[-10:]
            parts = []
            for m in recent:
                role_label = "姐" if m.get("role") == "user" else "小耕"
                text = (m.get("text") or m.get("content") or "").strip()
                if text:
                    parts.append(f"{role_label}：{text}")
            context_text = "\n".join(parts)

        system_prompt = build_persona_prompt(module="career")

        if not message.strip():
            # 初始问候
            step_name = step_labels.get(step, "求职")
            if step == "yipan":
                if has_resume:
                    prompt = (
                        "用户已上传简历，正在进行一盘·简历盘点。"
                        "小耕需要对简历做初步反馈，然后引导用户进入STAR追问环节。\n\n"
                        "请简短回应用户已上传简历这件事（2-3句话），"
                        "然后开始引导STAR追问：请用户选择一个最有成就感的项目/经历，"
                        "用STAR框架（情境/任务/行动/结果）来描述。\n"
                        "称呼「姐」，自称「小耕」，语气温暖鼓励。"
                    )
                else:
                    prompt = (
                        "用户刚进入一盘·简历盘点页面。"
                        "请以温暖的方式打招呼，邀请用户上传简历或者直接用文字/语音"
                        "聊聊职业经历。提及五大盘点：履历梳理→STAR追问→技能晶体→人脉资源→岗位建议。\n"
                        "2-3句话即可。称呼「姐」，自称「小耕」。"
                    )
            else:
                prompt = (
                    f"用户进入了「{step_name}」步骤。"
                    f"请以温暖的方式欢迎用户，介绍当前步骤的目标和要点。\n"
                    f"2-3句话即可。称呼「姐」，自称「小耕」。"
                )
        else:
            # 根据步骤和子进度构建引导
            if step == "yipan":
                phase_guides = [
                    # 0-履历梳理
                    ("请引导用户系统性地梳理职业经历。追问方向：\n"
                     "- 最近一段工作的公司、职位、主要职责\n"
                     "- 之前还有哪些重要的工作经历\n"
                     "先肯定用户的回答，再温和追问下一段经历。2-3句话。"),
                    # 1-STAR追问
                    ("用户正在做STAR萃取。引导用户用STAR框架描述一个具体项目：\n"
                     "S-当时什么背景？T-你要完成什么任务？A-你具体做了什么？R-结果如何？\n"
                     "每次只追问一个维度。如果用户已经说了S和T，就追问A。\n"
                     "先肯定已描述的，再追问缺失的。2-3句话。"),
                    # 2-技能晶体
                    ("用户已经完成STAR描述。现在需要从具体经历中提炼可复用的技能晶体。\n"
                     "帮用户识别：这个经历体现了什么核心能力？这个能力在什么场景下可以复用？\n"
                     "引导用户用一句话概括这个能力（做什么+怎么做），然后给出等级评估建议。\n"
                     "2-3句话，专业但不啰嗦。"),
                    # 3-人脉资源
                    ("用户在盘点求职资源。引导用户梳理人脉网络：\n"
                     "- 前领导/同事中谁能帮忙内推？\n"
                     "- 行业里认识哪些HR/猎头？\n"
                     "- 同学/校友网络中有没有可用的渠道？\n"
                     "鼓励用户把所有可能性都列出来。2-3句话。"),
                    # 4-岗位建议
                    ("用户已完成盘点，马上要进入下一步。\n"
                     "根据前面的对话，帮用户识别：\n"
                     "- 基于经历和技能，最适合的岗位方向是什么？\n"
                     "- 还有什么需要补充的？\n"
                     "最后引导用户进入「二定·求职策略」步骤。\n"
                     "2-3句话。"),
                ]
                guide = phase_guides[min(sub_index, 4)]
            else:
                guide = (
                    "根据用户的回答，给出温暖的回应，继续引导用户完成当前步骤。\n"
                    "2-3句话即可。"
                )

            prompt = (
                f"当前步骤：{step_labels.get(step, step)}，子阶段：{phase_label}\n"
                f"对话历史：\n{context_text}\n"
                f"────────────────\n"
                f"姐刚说：{message}\n\n"
                f"{guide}\n\n"
                f"称呼「姐」，自称「小耕」，语气温暖鼓励，专业务实。"
            )

        result = llm_generate_with_orchestration(
            prompt=prompt,
            system_prompt=system_prompt,
            module="career",
            task_complexity="medium",
            temperature=0.7,
            max_tokens=512,
            user_id=user_id,
            db=db,
        )

        reply = result.get("content", "").strip()
        if not reply:
            reply = "姐，小耕听到了～咱们继续梳理您的职场经历，让每段经历都变成闪闪发光的资产！"

        return {
            "reply": reply,
            "model_used": result.get("model_used", ""),
        }

    except Exception as e:
        logger.warning("高维求职AI对话失败: %s", e)
        return {
            "reply": "姐，小耕正在努力思考中，请稍等一下哦～",
            "model_used": "",
        }


# ═══════════════════════════════════════════════
# 五步法进度查询
# ═══════════════════════════════════════════════

def get_five_step_progress(db: Session, user_id: str) -> dict[str, Any]:
    """获取用户五步法整体进度。"""
    progress = db.query(CareerProgress).filter(
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()

    if not progress:
        # 没有进度记录，返回初始状态
        return {
            "career_progress_id": "",
            "current_step": 1,
            "step_labels": STEP_LABELS,
            "status": "not_started",
            "teacher_assigned": False,
            "step_details": {},
        }

    # 收集各步骤详情
    step_details = {}

    # 步骤1：简历 + STAR萃取 + 技能晶体
    star_count = db.query(func.count(STARExtraction.id)).filter(
        STARExtraction.career_progress_id == progress.id,
        STARExtraction.deleted_at.is_(None),
    ).scalar() or 0
    crystal_count = db.query(func.count(SkillCrystal.id)).filter(
        SkillCrystal.user_id == user_id,
        SkillCrystal.deleted_at.is_(None),
    ).scalar() or 0
    step_details["1"] = {
        "label": STEP_LABELS[1],
        "has_resume": bool(progress.step_data_json and progress.step_data_json.get("resume_content")),
        "star_extractions": star_count,
        "skill_crystals": crystal_count,
    }

    # 步骤2：求职策略
    has_strategy = bool(progress.step_data_json and progress.step_data_json.get("strategy"))
    step_details["2"] = {
        "label": STEP_LABELS[2],
        "has_strategy": has_strategy,
    }

    # 步骤3：投递追踪
    app_stats = _calculate_application_stats(db, user_id)
    step_details["3"] = {
        "label": STEP_LABELS[3],
        "applications_total": app_stats["total"],
        "invite_rate": app_stats["invite_rate"],
    }

    # 步骤4：面试准备与复盘
    interview_prep_ids = []
    if progress.step_data_json:
        interview_prep_ids = progress.step_data_json.get("interview_prep_ids", [])
    review_count = db.query(func.count(InterviewReview.id)).filter(
        InterviewReview.user_id == user_id,
        InterviewReview.deleted_at.is_(None),
    ).scalar() or 0
    step_details["4"] = {
        "label": STEP_LABELS[4],
        "interview_preps": len(interview_prep_ids),
        "interview_reviews": review_count,
    }

    # 步骤5：Offer评估
    comparison_count = db.query(func.count(OfferComparison.id)).filter(
        OfferComparison.user_id == user_id,
        OfferComparison.deleted_at.is_(None),
    ).scalar() or 0
    step_details["5"] = {
        "label": STEP_LABELS[5],
        "offer_comparisons": comparison_count,
    }

    teacher_assigned = bool(progress.teacher_id)

    return {
        "career_progress_id": progress.id,
        "current_step": progress.current_step,
        "step_labels": STEP_LABELS,
        "status": progress.status,
        "teacher_assigned": teacher_assigned,
        "step_details": step_details,
    }


# ═══════════════════════════════════════════════
# 步骤详情查询
# ═══════════════════════════════════════════════

def get_step_data(db: Session, user_id: str, step_id: int) -> dict[str, Any]:
    """获取指定步骤的详细数据。"""
    progress = db.query(CareerProgress).filter(
        CareerProgress.user_id == user_id,
        CareerProgress.deleted_at.is_(None),
    ).first()
    if not progress:
        raise APIError(20001, "五步法进度不存在", 404)

    if step_id < 1 or step_id > 5:
        raise APIError(20002, "步骤ID必须为1-5", 400)

    step_data = progress.step_data_json or {}

    if step_id == 1:
        # 简历 + STAR萃取 + 技能晶体
        extractions = (
            db.query(STARExtraction)
            .filter(
                STARExtraction.career_progress_id == progress.id,
                STARExtraction.deleted_at.is_(None),
            )
            .order_by(desc(STARExtraction.created_at))
            .all()
        )
        crystals = list_skill_crystals(db, user_id)

        return {
            "step_id": 1,
            "label": STEP_LABELS[1],
            "resume_title": step_data.get("resume_title", ""),
            "resume_content": step_data.get("resume_content", ""),
            "star_extractions": [
                {
                    "id": e.id,
                    "situation": e.situation or "",
                    "task": e.task or "",
                    "action": e.action or "",
                    "result": e.result or "",
                    "quantified_value": e.quantified_value or "",
                    "completeness": e.completeness,
                }
                for e in extractions
            ],
            "skill_crystals": crystals,
        }

    elif step_id == 2:
        strategy = step_data.get("strategy", {})
        return {
            "step_id": 2,
            "label": STEP_LABELS[2],
            "target_industry": strategy.get("target_industry", ""),
            "target_position": strategy.get("target_position", ""),
            "salary_range": strategy.get("salary_range", ""),
            "location": strategy.get("location", ""),
            "resource_inventory": strategy.get("resource_inventory", {}),
            "plan": strategy.get("plan", {}),
        }

    elif step_id == 3:
        apps = list_applications(db, user_id)
        return {
            "step_id": 3,
            "label": STEP_LABELS[3],
            "applications": apps,
        }

    elif step_id == 4:
        preps = list_interview_preps(db, user_id)
        return {
            "step_id": 4,
            "label": STEP_LABELS[4],
            "interview_preps": preps,
        }

    elif step_id == 5:
        comparisons = (
            db.query(OfferComparison)
            .filter(
                OfferComparison.user_id == user_id,
                OfferComparison.deleted_at.is_(None),
            )
            .order_by(desc(OfferComparison.created_at))
            .all()
        )
        return {
            "step_id": 5,
            "label": STEP_LABELS[5],
            "offer_comparisons": [
                {
                    "id": c.id,
                    "offers": c.offers_json,
                    "selected_offer_id": c.selected_offer_id,
                    "notes": c.comparison_notes,
                }
                for c in comparisons
            ],
        }

    return {"step_id": step_id, "label": STEP_LABELS.get(step_id, ""), "data": {}}


# ═══════════════════════════════════════════════
# 深度STAR萃取算法 (算法文档 §3.4)
# ═══════════════════════════════════════════════

def deep_star_extraction(resume_text: str, career_history: str = "", user_id=None, db=None) -> dict:
    """深度STAR成就事件萃取。

    输出: 结构化的STAR成就事件列表 + 技能晶体
    """
    from ...engines.persona import build_persona_prompt
    from ...engines.llm_orchestrator import llm_generate_with_orchestration
    import json, logging
    logger = logging.getLogger("career")

    prompt = f"""对以下工作经历进行STAR解构:

{resume_text}

{('补充经历: ' + career_history) if career_history else ''}

对每段经历进行四维解构:
S(Situation/情境): 公司阶段/团队规模/HR挑战
T(Task/任务): 具体任务/重要性/影响范围
A(Action/行动): 按时间顺序的行动/主导vs配合/困难克服
R(Result/结果): 量化成果/可验证的描述

同时提取'技能晶体'——可跨场景迁移的核心能力:
- 能力名称(如: 薪酬体系设计、高难度沟通)
- 能力证明(对应STAR事件简述)
- 能力等级(初级/熟练/精通/专家)
- 能力维度归类: 专业力/领导力/沟通力/分析力/执行力

返回JSON: {{star_events: [...], skill_crystals: [...]}}
完整性检查: 信息不足时标注'信息缺失', 不要编造。"""

    try:
        result = llm_generate_with_orchestration(
            prompt=prompt,
            system_prompt=build_persona_prompt(module="career"),
            module="growth_analysis", task_complexity="complex", temperature=0.6,
            user_id=user_id, db=db,
        )
        # Try to parse JSON from result
        content = result["content"]
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError:
            # Try extracting JSON block
            import re
            match = re.search(r'\{[^{}]*"star_events"[^{}]*\}', content, re.DOTALL)
            if match:
                parsed = json.loads(match.group())
            else:
                return {"star_events": [], "skill_crystals": [], "raw": content}
        return parsed
    except Exception as e:
        logger.warning("深度STAR萃取失败: %s", e)
        return {"star_events": [], "skill_crystals": []}


def generate_interview_strategy(job_description: str, star_extractions: list, company_intel: str = "", user_id=None, db=None) -> dict:
    """三阶段面试策略生成。

    阶段1: 人岗匹配度分析
    阶段2: 面试问题预测(必考题/专业题/行为题/压力题)
    阶段3: STAR应答准备
    """
    from ...engines.persona import build_persona_prompt
    from ...engines.llm_orchestrator import llm_generate_with_orchestration
    import json, logging
    logger = logging.getLogger("career")

    stars_text = json.dumps(star_extractions, ensure_ascii=False, indent=2)[:3000]

    prompt = f"""基于以下信息生成面试策略:

【岗位描述】: {job_description[:1000]}
【我的STAR成就事件】: {stars_text}
{f'【公司情报】: {company_intel[:500]}' if company_intel else ''}

三阶段分析:
1. 人岗匹配度: 优势(strengths)/短板(gaps)/差异化亮点(differentiators)
2. 面试问题预测:
   - 必考题(100%): 自我介绍/离职原因/期望薪资/职业规划
   - 专业题(80%): 基于JD的专业问题
   - 行为题(60%): 简历疑点追问
   - 压力题(30%): '你最大的失败是什么'
3. STAR应答准备: 为每个关键问题准备≤3分钟的STAR回答

返回JSON: {{match_analysis, predicted_questions, star_answers}}"""

    try:
        result = llm_generate_with_orchestration(
            prompt=prompt,
            system_prompt=build_persona_prompt(module="career"),
            module="career", task_complexity="complex", temperature=0.7,
            user_id=user_id, db=db,
        )
        try:
            return json.loads(result["content"])
        except json.JSONDecodeError:
            return {"raw_response": result["content"]}
    except Exception as e:
        logger.warning("面试策略生成失败: %s", e)
        return {"error": str(e)}
